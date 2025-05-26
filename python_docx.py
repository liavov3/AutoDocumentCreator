import pandas as pd
from docx import Document
from docx.text.run import WD_BREAK
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT,WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

### For all styles serch: 'Understanding Styles — python-docx 1.1.0 documentation'
def set_rtl(paragraph):
    """Apply RTL direction to a paragraph using XML hacks."""
    p = paragraph._p  # Access the internal XML <w:p> element
    pPr = p.get_or_add_pPr()

    # Create or get <w:bidi>
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)


def add_title(doc, title: str, language: str = None):
    if language == 'Hebrew':
        t = doc.add_heading(title, level=1)
        t.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Align to right

        set_rtl(t)

        # Set font for the heading (via run)
        run = t.runs[0]
        run.font.name = 'David'
        run.font.size = Pt(22)
    else:
        doc.add_heading(title, 0).alignment = 0


def add_p(doc, language: str = None):
    p = doc.add_paragraph()
    run = p.add_run()
    font = run.font
    font.name = 'David'  # Or another Hebrew-compatible font
    font.size = Pt(16)
    if language == 'Hebrew':
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    else:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return p


def add_text(p, text: str, style=None):
    if style:
        p.add_run(text).bold = True
    else:
        p.add_run(text)


def add_bullet_list(doc, text:str,  language: str = None):
    if language == 'Hebrew':
        p = doc.add_paragraph(text, style="List Bullet")
        p.alignment = 2
    else:
        doc.add_paragraph(text, style="List Bullet").alignment = 0


def create_rtl_table(doc, rows, cols):
    table = doc.add_table(rows=rows, cols=cols)

    tblPr = table._element.xpath('.//w:tblPr')[0]

    bidiVisual = OxmlElement('w:bidiVisual')
    tblPr.append(bidiVisual)

    return doc

def create_table(doc, rows: int, cols: int, data):
    table = doc.add_table(rows=1, cols=cols, style='Colorful Shading')
    df = pd.DataFrame(data)
    table_header = list(df.iloc[0])
    for i in range(cols):
        table.rows[0].cells[i].text = str(table_header[i])
    
    for i in range(rows):
        cells = table.add_row().cells
        for j in range(cols):
            cells[j].text = str(df.iloc[i+1][j])


def add_image(p, path: str):
    run = p.add_run()
    # first picture
    run.add_picture(path)
    # a single line‐break (not a page‐break!)
    run.add_break(WD_BREAK.LINE)
